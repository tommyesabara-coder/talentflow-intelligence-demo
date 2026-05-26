from __future__ import annotations

import cgi
import hashlib
import io
import json
import mimetypes
import os
import random
import re
import time
import urllib.error
import urllib.request
import base64
from http import HTTPStatus
from http.server import SimpleHTTPRequestHandler, ThreadingHTTPServer
from pathlib import Path
from typing import Any

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).parent
PORT = int(os.getenv("PORT", "8765"))
CURRENT_YEAR = 2026
CURRENT_MONTH = 5
OPENAI_PRIMARY_RESUME_MODEL = "gpt-4.1-mini"
OPENAI_FALLBACK_RESUME_MODEL = "gpt-4.1-nano"
GEMINI_DEFAULT_RESUME_MODEL = "gemini-2.5-flash"
TOGETHER_DEFAULT_MODEL = "openai/gpt-oss-20b"
TOGETHER_DEFAULT_BASE_URL = "https://api.together.ai/v1"
OPENROUTER_DEFAULT_MODEL = "google/gemini-2.5-flash"
OPENROUTER_DEFAULT_BASE_URL = "https://openrouter.ai/api/v1"
OLLAMA_DEFAULT_MODEL = "llama3.1:8b"
OLLAMA_DEFAULT_BASE_URL = "http://127.0.0.1:11434"

SUPPORTED_EXTENSIONS = {".pdf", ".doc", ".docx"}
PARSE_CACHE: dict[str, dict[str, Any]] = {}
PARSE_CACHE_MAX_ITEMS = 120
RUNTIME_CONFIG: dict[str, str] = {}
RUNTIME_CONFIG_FILE = ROOT / ".runtime-ai-config.json"
SKILL_KEYWORDS = sorted(
    {
        "administration",
        "analytics",
        "assessment",
        "ats",
        "audit",
        "benefit",
        "bpjs",
        "budgeting",
        "business partnering",
        "candidate experience",
        "coaching",
        "compensation",
        "competency matrix",
        "dashboard",
        "data analysis",
        "employee engagement",
        "employee relations",
        "employer branding",
        "excel",
        "ga",
        "google workspace",
        "headcount",
        "hr analytics",
        "hr business partner",
        "hr generalist",
        "hr operation",
        "hr operations",
        "hris",
        "inventory",
        "industrial relation",
        "interview",
        "job evaluation",
        "labor law",
        "lean",
        "logistic",
        "kpi",
        "learning development",
        "manpower planning",
        "microsoft office",
        "od",
        "onboarding",
        "payroll",
        "people analytics",
        "people data",
        "performance management",
        "power bi",
        "powerpoint",
        "procurement",
        "production planning",
        "recruitment",
        "safety",
        "sap",
        "sourcing",
        "sql",
        "stakeholder management",
        "stock opname",
        "succession planning",
        "supply chain",
        "talent acquisition",
        "talent management",
        "training",
        "union relation",
        "warehouse",
        "workforce planning",
    }
)


def local_env() -> dict[str, str]:
    values: dict[str, str] = {}
    for env_file in (ROOT / ".env", ROOT / "openai.env"):
        if not env_file.exists():
            continue
        for line in env_file.read_text(encoding="utf-8").splitlines():
            cleaned = line.strip()
            if not cleaned or cleaned.startswith("#") or "=" not in cleaned:
                continue
            key, value = cleaned.split("=", 1)
            values[key.strip().lstrip("\ufeff")] = value.strip().strip('"').strip("'")
    return values


def load_runtime_config() -> dict[str, str]:
    if RUNTIME_CONFIG:
        return RUNTIME_CONFIG
    if not RUNTIME_CONFIG_FILE.exists():
        return RUNTIME_CONFIG
    try:
        payload = json.loads(RUNTIME_CONFIG_FILE.read_text(encoding="utf-8"))
        if isinstance(payload, dict):
            for key, value in payload.items():
                if isinstance(key, str) and isinstance(value, str):
                    RUNTIME_CONFIG[key] = value
    except Exception:
        return RUNTIME_CONFIG
    return RUNTIME_CONFIG


def save_runtime_config() -> None:
    safe_payload = {key: value for key, value in RUNTIME_CONFIG.items() if isinstance(value, str)}
    RUNTIME_CONFIG_FILE.write_text(json.dumps(safe_payload, ensure_ascii=False, indent=2), encoding="utf-8")


def get_config(name: str, default: str = "") -> str:
    return load_runtime_config().get(name) or os.getenv(name) or local_env().get(name, default)


def urlopen_direct(request: urllib.request.Request, timeout: int | float = 75):
    """Open external AI/API requests without inheriting broken local proxy env vars."""
    opener = urllib.request.build_opener(urllib.request.ProxyHandler({}))
    return opener.open(request, timeout=timeout)


def get_openai_api_key() -> str:
    value = get_config("OPENAI_API_KEY").strip()
    if value.upper().startswith("OPENAI_API_KEY="):
        value = value.split("=", 1)[1].strip()
    placeholders = {
        "",
        "isi_api_key_openai_di_sini",
        "your_openai_api_key_here",
        "sk-...",
        "sk-proj-...",
    }
    return "" if value.lower() in placeholders else value


def get_gemini_api_key() -> str:
    value = get_config("GEMINI_API_KEY").strip()
    if value.upper().startswith("GEMINI_API_KEY="):
        value = value.split("=", 1)[1].strip()
    placeholders = {"", "isi_key_gemini_anda", "your_gemini_api_key_here"}
    return "" if value.lower() in placeholders else value


def get_together_api_key() -> str:
    value = get_config("TOGETHER_API_KEY").strip()
    if value.upper().startswith("TOGETHER_API_KEY="):
        value = value.split("=", 1)[1].strip()
    placeholders = {"", "isi_key_together_anda", "your_together_api_key_here", "tgp_v1_..."}
    return "" if value.lower() in placeholders else value


def get_openrouter_api_key() -> str:
    value = get_config("OPENROUTER_API_KEY").strip()
    if value.upper().startswith("OPENROUTER_API_KEY="):
        value = value.split("=", 1)[1].strip()
    placeholders = {"", "isi_key_openrouter_anda", "your_openrouter_api_key_here", "sk-or-v1-..."}
    return "" if value.lower() in placeholders else value


def get_resume_models() -> list[str]:
    configured = get_config("OPENAI_RESUME_MODEL", OPENAI_PRIMARY_RESUME_MODEL)
    fallback = get_config("OPENAI_RESUME_FALLBACK_MODEL", OPENAI_FALLBACK_RESUME_MODEL)
    return dedupe([configured, fallback])


def get_gemini_resume_model() -> str:
    return get_config("GEMINI_RESUME_MODEL", GEMINI_DEFAULT_RESUME_MODEL)


def get_together_model() -> str:
    model = get_config("TOGETHER_MODEL", TOGETHER_DEFAULT_MODEL).strip()
    return TOGETHER_DEFAULT_MODEL if model.lower() in {"", "together"} else model


def get_together_base_url() -> str:
    return get_config("TOGETHER_BASE_URL", TOGETHER_DEFAULT_BASE_URL).strip().rstrip("/") or TOGETHER_DEFAULT_BASE_URL


def get_openrouter_model() -> str:
    return get_config("OPENROUTER_MODEL", OPENROUTER_DEFAULT_MODEL).strip() or OPENROUTER_DEFAULT_MODEL


def get_openrouter_base_url() -> str:
    return get_config("OPENROUTER_BASE_URL", OPENROUTER_DEFAULT_BASE_URL).strip().rstrip("/") or OPENROUTER_DEFAULT_BASE_URL


def get_ollama_base_url() -> str:
    return get_config("OLLAMA_BASE_URL", OLLAMA_DEFAULT_BASE_URL).strip().rstrip("/") or OLLAMA_DEFAULT_BASE_URL


def get_ollama_model() -> str:
    return get_config("OLLAMA_RESUME_MODEL", OLLAMA_DEFAULT_MODEL).strip() or OLLAMA_DEFAULT_MODEL


def get_ai_engine_strategy() -> str:
    strategy = get_config("AI_ENGINE_STRATEGY", "random").strip().lower()
    return strategy if strategy in {"random", "ordered"} else "random"


def local_only_parser_enabled() -> bool:
    value = get_config("LOCAL_ONLY_PARSER", "true").strip().lower()
    return value not in {"0", "false", "no", "off"}


def hybrid_fast_parser_enabled() -> bool:
    value = get_config("HYBRID_FAST_PARSER", "true").strip().lower()
    return value not in {"0", "false", "no", "off"}


def parse_confidence_threshold() -> float:
    try:
        return min(max(float(get_config("HYBRID_CONFIDENCE_THRESHOLD", "0.78")), 0.0), 1.0)
    except ValueError:
        return 0.78


def field_confidence_threshold() -> float:
    try:
        return min(max(float(get_config("HYBRID_FIELD_CONFIDENCE_THRESHOLD", "0.64")), 0.0), 1.0)
    except ValueError:
        return 0.64


def ollama_parser_enabled() -> bool:
    value = get_config("OLLAMA_ENABLED", "true").strip().lower()
    return value not in {"0", "false", "no", "off"}


def get_ollama_status(timeout: float = 0.8) -> dict[str, Any]:
    if not ollama_parser_enabled():
        return {"available": False, "models": [], "error": "disabled"}
    try:
        request = urllib.request.Request(f"{get_ollama_base_url()}/api/tags", method="GET")
        with urlopen_direct(request, timeout=timeout) as response:  # noqa: S310
            payload = json.loads(response.read().decode("utf-8"))
        models = [
            str(item.get("name") or "")
            for item in payload.get("models", [])
            if isinstance(item, dict) and item.get("name")
        ]
        return {"available": True, "models": models, "error": ""}
    except Exception as exc:  # noqa: BLE001
        return {"available": False, "models": [], "error": str(exc)[:160]}


def configured_ai_engines() -> list[str]:
    preferred = [
        engine.strip().lower()
        for engine in get_config("AI_ENGINE_PROVIDERS", "gemini,openai,openrouter,together").split(",")
        if engine.strip()
    ]
    engines: list[str] = []
    for engine in preferred:
        if engine == "gemini" and get_gemini_api_key():
            engines.append("gemini")
        if engine == "openai" and get_openai_api_key():
            engines.append("openai")
        if engine == "together" and get_together_api_key():
            engines.append("together")
        if engine == "openrouter" and get_openrouter_api_key():
            engines.append("openrouter")
    return dedupe(engines)


def runtime_ai_key_configured() -> bool:
    return any(
        load_runtime_config().get(key)
        for key in ("GEMINI_API_KEY", "OPENAI_API_KEY", "TOGETHER_API_KEY", "OPENROUTER_API_KEY")
    )


def detect_api_key_provider(api_key: str) -> str:
    lowered = api_key.strip().lower()
    if lowered.startswith(("sk-or-", "sk-or-v1-")):
        return "openrouter"
    if lowered.startswith(("sk-", "sk-proj-")):
        return "openai"
    if lowered.startswith("tgp_") or lowered.startswith("together_"):
        return "together"
    if api_key.strip().startswith("AIza"):
        return "gemini"
    return ""


def validate_runtime_api_key_shape(provider: str, api_key: str) -> str:
    compact = api_key.strip()
    if provider == "openai" and len(compact) < 18:
        return "API key terlihat terlalu pendek. Paste full key asli, bukan key yang terpotong/termasking."
    if provider == "gemini" and len(compact) < 30:
        return "Gemini API key terlihat terlalu pendek. Paste full key dari Google AI Studio."
    if provider == "together" and len(compact) < 30:
        return "Together API key terlihat terlalu pendek. Paste full key dari dashboard Together."
    if provider == "openrouter" and len(compact) < 30:
        return "OpenRouter API key terlihat terlalu pendek. Paste full key dari dashboard OpenRouter."
    if "*" in compact or "…" in compact or "..." in compact:
        return "API key terlihat masih masking/terpotong. Paste full key asli, bukan tampilan yang sudah disensor."
    return ""


def select_ai_engines() -> list[str]:
    engines = configured_ai_engines()
    if get_ai_engine_strategy() == "random":
        random.shuffle(engines)
    return engines


def build_ai_status() -> dict[str, Any]:
    if local_only_parser_enabled():
        ollama_status = get_ollama_status()
        ollama_ready = bool(ollama_status["available"])
        configured_model = get_ollama_model()
        model_available = configured_model in ollama_status.get("models", [])
        return {
            "aiConfigured": ollama_ready,
            "openaiConfigured": bool(get_openai_api_key()),
            "geminiConfigured": bool(get_gemini_api_key()),
            "togetherConfigured": bool(get_together_api_key()),
            "openrouterConfigured": bool(get_openrouter_api_key()),
            "ollamaConfigured": ollama_ready,
            "mode": f"Ollama Local AI â€¢ {configured_model}" if ollama_ready else "Local Max Parser",
            "primaryModel": configured_model if ollama_ready else "Local rules + evidence scoring",
            "fallbackModels": ["Local Max Parser"] if ollama_ready else [],
            "engineStrategy": "local-first",
            "availableEngines": ["ollama", "local-rules"] if ollama_ready else ["local-rules"],
            "keySource": "local Ollama service" if ollama_ready else "",
            "message": (
                "Ollama aktif: parser memakai AI lokal tanpa API berbayar."
                if ollama_ready and model_available
                else "Ollama service aktif, tetapi model default belum terlihat. Jalankan ollama pull model yang dipilih."
                if ollama_ready
                else "Parser lokal aktif. Ollama belum running di 127.0.0.1:11434."
            ),
            "ollamaModels": ollama_status.get("models", []),
            "ollamaModelAvailable": model_available,
        }

    openai_key = get_openai_api_key()
    gemini_key = get_gemini_api_key()
    together_key = get_together_api_key()
    openrouter_key = get_openrouter_api_key()
    active_engines = configured_ai_engines()
    providers = []
    if "openai" in active_engines and openai_key:
        providers.append(f"OpenAI {get_resume_models()[0]}")
    if "gemini" in active_engines and gemini_key:
        providers.append(f"Gemini {get_gemini_resume_model()}")
    if "together" in active_engines and together_key:
        providers.append(f"Together {get_together_model()}")
    if "openrouter" in active_engines and openrouter_key:
        providers.append(f"OpenRouter {get_openrouter_model()}")
    ai_configured = bool(providers)
    return {
        "aiConfigured": ai_configured,
        "openaiConfigured": bool(openai_key),
        "geminiConfigured": bool(gemini_key),
        "togetherConfigured": bool(together_key),
        "openrouterConfigured": bool(openrouter_key),
        "mode": " + ".join(providers) if providers else "Local Smart Parser",
        "primaryModel": providers[0] if providers else get_resume_models()[0],
        "fallbackModels": providers[1:],
        "engineStrategy": get_ai_engine_strategy(),
        "availableEngines": active_engines,
        "keySource": "Input API Key" if runtime_ai_key_configured() else "server environment/.env" if ai_configured else "",
        "runtimeApiKeyConfigured": runtime_ai_key_configured(),
        "message": "AI parser aktif" if ai_configured else "AI parser belum aktif: pasang API key AI dari tombol Input API Key.",
        "hybridFastParser": hybrid_fast_parser_enabled(),
        "hybridConfidenceThreshold": parse_confidence_threshold(),
        "cacheItems": len(PARSE_CACHE),
    }

SECTION_WORDS = {
    "curriculum vitae",
    "resume",
    "profile",
    "profil",
    "summary",
    "ringkasan",
    "work experience",
    "pengalaman kerja",
    "experience",
    "employment history",
    "education",
    "pendidikan",
    "skills",
    "skill",
    "keahlian",
    "contact",
    "kontak",
    "personal data",
    "data pribadi",
    "certification",
    "certifications",
    "sertifikasi",
}

JOB_WORDS = {
    "analyst",
    "assistant",
    "business",
    "chief",
    "consultant",
    "coordinator",
    "director",
    "engineer",
    "executive",
    "general",
    "head",
    "hr",
    "human",
    "intern",
    "lead",
    "leader",
    "manager",
    "marketing",
    "officer",
    "operation",
    "operations",
    "partner",
    "people",
    "recruiter",
    "resources",
    "specialist",
    "staff",
    "supervisor",
    "talent",
}

MONTHS = {
    "jan": 1,
    "january": 1,
    "januari": 1,
    "feb": 2,
    "february": 2,
    "februari": 2,
    "mar": 3,
    "march": 3,
    "maret": 3,
    "apr": 4,
    "april": 4,
    "may": 5,
    "mei": 5,
    "jun": 6,
    "june": 6,
    "juni": 6,
    "jul": 7,
    "july": 7,
    "juli": 7,
    "aug": 8,
    "august": 8,
    "agustus": 8,
    "sep": 9,
    "sept": 9,
    "september": 9,
    "oct": 10,
    "october": 10,
    "okt": 10,
    "oktober": 10,
    "nov": 11,
    "november": 11,
    "dec": 12,
    "december": 12,
    "des": 12,
    "desember": 12,
}


class RecruitFlowHandler(SimpleHTTPRequestHandler):
    def translate_path(self, path: str) -> str:
        requested = super().translate_path(path)
        relative = Path(requested).relative_to(Path.cwd())
        return str(ROOT / relative)

    def end_headers(self) -> None:
        self.send_header("Access-Control-Allow-Origin", "*")
        self.send_header("Access-Control-Allow-Methods", "GET, POST, OPTIONS")
        self.send_header("Access-Control-Allow-Headers", "Content-Type")
        super().end_headers()

    def do_OPTIONS(self) -> None:
        self.send_response(HTTPStatus.NO_CONTENT)
        self.end_headers()

    def do_GET(self) -> None:
        if self.path.split("?", 1)[0] == "/api/ai-status":
            self.send_json(build_ai_status())
            return
        super().do_GET()

    def do_POST(self) -> None:
        path = self.path.split("?", 1)[0]
        if path == "/api/runtime-ai-key":
            self.handle_runtime_ai_key()
            return
        if path == "/api/reanalyze-cv":
            self.handle_reanalyze_cv()
            return
        if path == "/api/generate-blueprint":
            self.handle_generate_blueprint()
            return
        if path == "/api/score-candidate-fit":
            self.handle_score_candidate_fit()
            return

        if path != "/api/parse-cvs":
            self.send_error(HTTPStatus.NOT_FOUND, "Unknown endpoint")
            return

        try:
            form = cgi.FieldStorage(
                fp=self.rfile,
                headers=self.headers,
                environ={
                    "REQUEST_METHOD": "POST",
                    "CONTENT_TYPE": self.headers.get("Content-Type", ""),
                },
            )
            files = form["files"] if "files" in form else []
            if not isinstance(files, list):
                files = [files]

            parsed: list[dict[str, Any]] = []
            errors: list[dict[str, str]] = []

            for item in files:
                filename = item.filename or "unknown"
                extension = Path(filename).suffix.lower()

                if extension not in SUPPORTED_EXTENSIONS:
                    errors.append({"file": filename, "error": "Format belum didukung"})
                    continue

                raw = item.file.read()
                try:
                    file_hash = resume_file_hash(raw)
                    cached_candidate = get_cached_resume(file_hash)
                    if cached_candidate:
                        cached_candidate["fileName"] = filename
                        cached_candidate["cacheHit"] = True
                        cached_candidate["parserMode"] = f"{cached_candidate.get('parserMode', 'Hybrid Parser')} â€¢ cache"
                        parsed.append(cached_candidate)
                        continue
                    text, extraction_meta = extract_text(raw, extension)
                    candidate = parse_resume(filename, text, extraction_meta, raw, extension)
                    candidate["fileHash"] = file_hash
                    set_cached_resume(file_hash, candidate)
                    parsed.append(candidate)
                except Exception as exc:  # noqa: BLE001
                    errors.append({"file": filename, "error": f"Gagal membaca file: {exc}"})

            self.send_json({"candidates": parsed, "errors": errors})
        except Exception as exc:  # noqa: BLE001
            self.send_json({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def handle_reanalyze_cv(self) -> None:
        try:
            content_length = int(self.headers.get("Content-Length", "0") or 0)
            payload = json.loads(self.rfile.read(content_length).decode("utf-8") or "{}")
            filename = clean_scalar(payload.get("fileName")) or "candidate-cv.txt"
            raw_text = payload.get("cvText") or payload.get("text") or ""
            text = clean_extracted_text(raw_text if isinstance(raw_text, str) else str(raw_text))
            if len(text.strip()) < 40:
                self.send_json(
                    {
                        "error": "Teks CV tersimpan terlalu pendek. Upload ulang CV asli berbasis teks, atau lakukan OCR jika PDF berupa scan/gambar.",
                    },
                    status=HTTPStatus.BAD_REQUEST,
                )
                return

            extraction_meta = {
                "sourceType": "Stored CV Text",
                "textLength": len(text),
                "warning": "Re-analysis memakai teks CV yang tersimpan di aplikasi. Jika sumbernya salah, upload ulang PDF/DOCX asli.",
            }
            candidate = parse_resume(filename, text, extraction_meta, b"", ".txt")
            candidate["reanalysisSource"] = "storedCvText"
            self.send_json({"candidate": candidate})
        except json.JSONDecodeError:
            self.send_json({"error": "Payload re-analysis bukan JSON valid."}, status=HTTPStatus.BAD_REQUEST)
        except Exception as exc:  # noqa: BLE001
            self.send_json({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def read_json_body(self) -> dict[str, Any]:
        content_length = int(self.headers.get("Content-Length", "0") or 0)
        return json.loads(self.rfile.read(content_length).decode("utf-8") or "{}")

    def handle_runtime_ai_key(self) -> None:
        try:
            payload = self.read_json_body()
            api_key = clean_scalar(payload.get("apiKey"))
            model = clean_scalar(payload.get("model"))

            provider_map = {
                "gemini": ("GEMINI_API_KEY", "GEMINI_RESUME_MODEL", GEMINI_DEFAULT_RESUME_MODEL),
                "openai": ("OPENAI_API_KEY", "OPENAI_RESUME_MODEL", OPENAI_PRIMARY_RESUME_MODEL),
                "together": ("TOGETHER_API_KEY", "TOGETHER_MODEL", TOGETHER_DEFAULT_MODEL),
                "openrouter": ("OPENROUTER_API_KEY", "OPENROUTER_MODEL", OPENROUTER_DEFAULT_MODEL),
            }
            requested_provider = clean_scalar(payload.get("provider")).lower() if "provider" in payload else ""
            key_probe = api_key.strip()
            key_probe_lower = key_probe.lower()
            auto_provider = ""
            if key_probe_lower.startswith(("sk-or-", "sk-or-v1-")):
                auto_provider = "openrouter"
            elif key_probe_lower.startswith("sk-"):
                auto_provider = "openai"
            elif key_probe_lower.startswith(("tgp_", "together_")):
                auto_provider = "together"
            elif key_probe.startswith("AIza"):
                auto_provider = "gemini"
            provider = requested_provider if requested_provider in provider_map else auto_provider

            if len(api_key) < 12:
                self.send_json({"error": "API key terlalu pendek atau kosong."}, status=HTTPStatus.BAD_REQUEST)
                return
            if provider not in provider_map:
                self.send_json(
                    {
                        "error": "Format API key belum dikenali otomatis. Saat ini sistem mengenali Gemini (AIza...), OpenAI (sk-...), OpenRouter (sk-or...), dan Together (tgp_...).",
                    },
                    status=HTTPStatus.BAD_REQUEST,
                )
                return
            shape_error = validate_runtime_api_key_shape(provider, api_key)
            if shape_error:
                self.send_json({"error": shape_error}, status=HTTPStatus.BAD_REQUEST)
                return

            key_name, model_name, default_model = provider_map[provider]
            RUNTIME_CONFIG[key_name] = api_key
            RUNTIME_CONFIG["AI_ENGINE_PROVIDERS"] = provider
            RUNTIME_CONFIG["AI_ENGINE_STRATEGY"] = "ordered"
            RUNTIME_CONFIG["LOCAL_ONLY_PARSER"] = "false"
            RUNTIME_CONFIG["HYBRID_FAST_PARSER"] = "true"
            if model:
                RUNTIME_CONFIG[model_name] = model
            else:
                RUNTIME_CONFIG[model_name] = default_model
            save_runtime_config()

            last_four = api_key[-4:] if len(api_key) >= 4 else "****"
            self.send_json(
                {
                    "ok": True,
                    "provider": provider,
                    "maskedKey": f"â€¢â€¢â€¢â€¢ {last_four}",
                    "status": build_ai_status(),
                }
            )
        except json.JSONDecodeError:
            self.send_json({"error": "Payload API key bukan JSON valid."}, status=HTTPStatus.BAD_REQUEST)
        except Exception as exc:  # noqa: BLE001
            self.send_json({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def handle_generate_blueprint(self) -> None:
        try:
            if local_only_parser_enabled():
                self.send_json({"error": "Mode lokal aktif: Jobdesc dibuat dari input user tanpa generate AI berbayar."}, status=HTTPStatus.BAD_REQUEST)
                return
            payload = self.read_json_body()
            brief = payload.get("brief") if isinstance(payload.get("brief"), dict) else payload
            blueprint = generate_blueprint_with_ai(brief)
            if not blueprint:
                self.send_json({"error": "Multi-engine AI belum berhasil membuat Job Blueprint."}, status=HTTPStatus.BAD_REQUEST)
                return
            self.send_json({"blueprint": blueprint})
        except json.JSONDecodeError:
            self.send_json({"error": "Payload Job Blueprint bukan JSON valid."}, status=HTTPStatus.BAD_REQUEST)
        except Exception as exc:  # noqa: BLE001
            self.send_json({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def handle_score_candidate_fit(self) -> None:
        try:
            if local_only_parser_enabled():
                self.send_json({"error": "Mode lokal aktif: scoring memakai inferensi lokal di aplikasi, bukan API berbayar."}, status=HTTPStatus.BAD_REQUEST)
                return
            payload = self.read_json_body()
            candidate = payload.get("candidate") if isinstance(payload.get("candidate"), dict) else {}
            blueprint = payload.get("blueprint") if isinstance(payload.get("blueprint"), dict) else {}
            score = score_candidate_fit_with_ai(candidate, blueprint)
            if not score:
                self.send_json({"error": "Multi-engine AI belum berhasil membuat fit score."}, status=HTTPStatus.BAD_REQUEST)
                return
            self.send_json({"fitScore": score})
        except json.JSONDecodeError:
            self.send_json({"error": "Payload scoring bukan JSON valid."}, status=HTTPStatus.BAD_REQUEST)
        except Exception as exc:  # noqa: BLE001
            self.send_json({"error": str(exc)}, status=HTTPStatus.BAD_REQUEST)

    def send_json(self, payload: dict[str, Any], status: HTTPStatus = HTTPStatus.OK) -> None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        self.send_response(status)
        self.send_header("Content-Type", "application/json; charset=utf-8")
        self.send_header("Content-Length", str(len(body)))
        self.end_headers()
        self.wfile.write(body)


def extract_text(raw: bytes, extension: str) -> tuple[str, dict[str, Any]]:
    if extension == ".pdf":
        reader = PdfReader(io.BytesIO(raw))
        pages = [clean_extracted_text(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(page for page in pages if page.strip())
        meta = {"sourceType": "PDF", "pageCount": len(reader.pages), "textLength": len(text)}
        if len(text.strip()) < 60:
            meta["lowTextExtraction"] = True
            meta["warning"] = "Teks PDF lokal sangat pendek; kemungkinan PDF hasil scan/gambar. Parser lokal butuh teks asli atau OCR."
        return text, meta

    if extension == ".docx":
        document = Document(io.BytesIO(raw))
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            parts.extend(extract_docx_table(table))
        for section in document.sections:
            parts.extend(paragraph.text for paragraph in section.header.paragraphs if paragraph.text.strip())
            parts.extend(paragraph.text for paragraph in section.footer.paragraphs if paragraph.text.strip())
        text = clean_extracted_text("\n".join(parts))
        meta = {"sourceType": "DOCX", "textLength": len(text), "tableCount": len(document.tables)}
        if len(text.strip()) < 40:
            raise ValueError("DOCX tidak berisi teks yang cukup untuk dianalisis.")
        return text, meta

    if extension == ".doc":
        text = extract_legacy_doc_text(raw)
        return text, {"sourceType": "DOC", "textLength": len(text), "warning": "Legacy DOC best-effort extraction"}

    raise ValueError("Unsupported extension")


def extract_docx_table(table: Any) -> list[str]:
    output: list[str] = []
    for row in table.rows:
        cells = [re.sub(r"\s+", " ", cell.text).strip() for cell in row.cells if cell.text.strip()]
        if cells:
            output.append(" | ".join(dedupe(cells)))
    return output


def clean_extracted_text(text: str) -> str:
    text = text.replace("\x00", " ").replace("\uf0b7", "â€¢")
    text = re.sub(r"(?<=\w)-\n(?=\w)", "", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_legacy_doc_text(raw: bytes) -> str:
    """Best-effort extraction for legacy .doc files."""
    text = raw.decode("utf-8", errors="ignore")
    if len(text.strip()) < 80:
        text = raw.decode("latin-1", errors="ignore")

    fragments = re.findall(r"[A-Za-z0-9@+.,:/()&_' -]{4,}", text)
    cleaned = "\n".join(
        re.sub(r"\s+", " ", fragment).strip()
        for fragment in fragments
        if re.search(r"[A-Za-z]{3,}", fragment)
    )

    if len(cleaned.strip()) < 40:
        raise ValueError("File DOC lama tidak bisa dibaca jelas. Mohon convert ke DOCX atau PDF.")

    return cleaned


def resume_file_hash(raw: bytes) -> str:
    return hashlib.sha256(raw).hexdigest()


def clone_jsonable(value: dict[str, Any]) -> dict[str, Any]:
    return json.loads(json.dumps(value, ensure_ascii=False))


def get_cached_resume(file_hash: str) -> dict[str, Any] | None:
    cached = PARSE_CACHE.get(file_hash)
    return clone_jsonable(cached) if cached else None


def set_cached_resume(file_hash: str, candidate: dict[str, Any]) -> None:
    if len(PARSE_CACHE) >= PARSE_CACHE_MAX_ITEMS:
        PARSE_CACHE.pop(next(iter(PARSE_CACHE)))
    PARSE_CACHE[file_hash] = clone_jsonable(candidate)


def needs_ai_resume_refinement(local_candidate: dict[str, Any], extraction_meta: dict[str, Any]) -> bool:
    if extraction_meta.get("lowTextExtraction"):
        return True
    confidence = float(local_candidate.get("parseConfidence") or 0)
    if confidence < parse_confidence_threshold():
        return True

    field_confidence = local_candidate.get("fieldConfidence") or {}
    important_fields = ["name", "experience", "education", "companies"]
    threshold = field_confidence_threshold()
    low_fields = [field for field in important_fields if float(field_confidence.get(field) or 0) < threshold]
    if low_fields:
        return True

    if not local_candidate.get("name") or not local_candidate.get("companies"):
        return True
    if safe_int(local_candidate.get("experience")) <= 0 and not local_candidate.get("rawExperience"):
        return True
    return False


def build_hybrid_resume_text(text: str, local_candidate: dict[str, Any], limit: int = 9000) -> str:
    """Compress CV text for selective AI fallback without losing the high-value evidence zones."""
    normalized = clean_extracted_text(text)
    lines = normalize_lines(normalized)
    work_zone = extract_work_zone(lines)

    high_value_lines: list[str] = []
    section_keywords = re.compile(
        r"(experience|pengalaman|employment|work|riwayat|education|pendidikan|skill|keahlian|certification|sertifikat|project|summary|profile|profil)",
        re.I,
    )
    for index, line in enumerate(lines):
        if index < 30 or section_keywords.search(line):
            high_value_lines.extend(lines[max(0, index - 2) : min(len(lines), index + 8)])

    draft = {
        key: local_candidate.get(key)
        for key in [
            "name",
            "email",
            "phone",
            "education",
            "companies",
            "jobTitles",
            "experience",
            "experienceMonths",
            "rawExperience",
            "skills",
            "fieldConfidence",
        ]
    }
    compressed = "\n".join(
        [
            "LOCAL PARSER DRAFT:",
            json.dumps(draft, ensure_ascii=False),
            "",
            "HIGH VALUE CV LINES:",
            "\n".join(dedupe(high_value_lines)[:260]),
            "",
            "WORK HISTORY ZONE:",
            "\n".join(work_zone[:220]),
            "",
            "CV START:",
            normalized[:3500],
        ]
    )
    return compressed[:limit]


def parse_resume(filename: str, text: str, extraction_meta: dict[str, Any], raw: bytes = b"", extension: str = "") -> dict[str, Any]:
    local_candidate = parse_resume_locally(filename, text, extraction_meta)
    if local_only_parser_enabled():
        ollama_errors: list[str] = []
        ollama_candidate = parse_resume_with_ollama(filename, text, local_candidate, ollama_errors)
        if ollama_candidate:
            return merge_ai_candidate(local_candidate, ollama_candidate)

        warning_prefix = ["Ollama belum aktif/berhasil; hasil memakai Local Max Parser."] if ollama_errors else []
        local_candidate["parserMode"] = "Local Max Parser"
        local_candidate["parseWarnings"] = dedupe(
            [
                *warning_prefix,
                *ollama_errors[:1],
                "Parser lokal aktif: tidak memakai API berbayar. Untuk PDF scan/gambar, hasil tetap bergantung pada teks yang bisa diekstrak.",
                *local_candidate.get("parseWarnings", []),
            ]
        )
        return local_candidate

    if hybrid_fast_parser_enabled() and not needs_ai_resume_refinement(local_candidate, extraction_meta):
        local_candidate["parserMode"] = "Hybrid Fast Parser â€¢ local high confidence"
        local_candidate["parseWarnings"] = dedupe(
            [
                "Fast mode: field utama terbaca confidence tinggi, AI tidak dipakai agar upload lebih cepat.",
                *local_candidate.get("parseWarnings", []),
            ]
        )[:8]
        return local_candidate

    ai_errors: list[str] = []
    ai_text = build_hybrid_resume_text(text, local_candidate) if hybrid_fast_parser_enabled() else text
    ai_candidate = parse_resume_with_multi_engine(filename, ai_text, local_candidate, raw, extension, ai_errors)

    if ai_candidate:
        merged = merge_ai_candidate(local_candidate, ai_candidate)
        if hybrid_fast_parser_enabled():
            merged["parserMode"] = f"Hybrid Accurate Parser â€¢ {merged.get('parserMode', 'AI fallback')}"
        return merged
    if not configured_ai_engines():
        local_candidate["parseWarnings"] = dedupe(
            [
                "AI parser belum aktif: belum ada API key yang dikenali. Hasil memakai Smart Parser lokal.",
                *local_candidate.get("parseWarnings", []),
            ]
        )
    elif ai_errors:
        local_candidate["parserMode"] = "Local Smart Parser â€¢ AI fallback"
        local_candidate["parseWarnings"] = dedupe(
            [
                f"AI parser belum berhasil: {ai_errors[0]} Hasil sementara memakai Smart Parser lokal.",
                *local_candidate.get("parseWarnings", []),
            ]
        )
    return local_candidate


def parse_resume_with_multi_engine(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    raw: bytes = b"",
    extension: str = "",
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    engines = select_ai_engines()
    if not engines:
        return None

    for engine in engines:
        if engine == "gemini":
            candidate = parse_resume_with_gemini(filename, text, local_candidate, raw, extension, errors)
        elif engine == "openai":
            candidate = parse_resume_with_ai(filename, text, local_candidate, raw, extension, errors)
        elif engine == "together":
            candidate = parse_resume_with_together(filename, text, local_candidate, errors)
        elif engine == "openrouter":
            candidate = parse_resume_with_openrouter(filename, text, local_candidate, errors)
        else:
            continue
        if candidate:
            candidate["engine"] = engine
            return candidate
    return None


def parse_resume_with_ollama(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    if not ollama_parser_enabled():
        return None
    try:
        parsed = call_ollama_json(build_ollama_resume_prompt(filename, text, local_candidate))
        if isinstance(parsed, dict):
            model = get_ollama_model()
            parsed["parserMode"] = f"Ollama Local Parser â€¢ {model}"
            parsed.setdefault("confidence", 0.76)
            parsed.setdefault("fieldConfidence", {})
            parsed.setdefault("evidence", [])
            parsed.setdefault("warnings", [])
            parsed.setdefault("skills", [])
            parsed.setdefault("companies", [])
            parsed.setdefault("jobTitles", [])
            return parsed
    except Exception as exc:  # noqa: BLE001
        if errors is not None:
            errors.append(format_ollama_error(exc))
    return None


def format_ollama_error(exc: Exception) -> str:
    message = str(exc)
    if "10061" in message or "actively refused" in message or "Connection refused" in message:
        return "Ollama belum running. Buka Ollama, lalu pastikan service aktif di 127.0.0.1:11434."
    if "model" in message.lower() and ("not found" in message.lower() or "pull" in message.lower()):
        return f"Model Ollama belum tersedia. Jalankan: ollama pull {get_ollama_model()}"
    return f"Ollama belum berhasil membaca CV: {message[:160]}."


def build_ollama_resume_prompt(filename: str, text: str, local_candidate: dict[str, Any]) -> str:
    return f"""
Anda adalah parser CV untuk recruiter Indonesia.
Tugas: perbaiki hasil parser lokal dengan membaca CV secara menyeluruh.

Aturan penting:
- Keluarkan JSON valid saja, tanpa markdown.
- Jangan mengarang data yang tidak ada di CV.
- Nama harus nama orang, bukan jabatan/perusahaan/judul section.
- Hitung pengalaman kerja dari rentang tanggal kerja jika ada; overlap dihitung satu kali.
- Jika pengalaman/skill kuat tersirat dari jabatan dan riwayat kerja, boleh masukkan sebagai inferensi profesional, tetapi evidence harus menjelaskan sumbernya.
- Gunakan Bahasa Indonesia untuk warnings, evidence.source, dan summary.

JSON wajib berisi field:
name, email, phone, age, education, companies, jobTitles, experience, experienceMonths, skills, summary,
confidence, fieldConfidence, evidence, warnings.

Draft parser lokal:
{json.dumps({k: local_candidate.get(k) for k in ['name', 'email', 'phone', 'age', 'education', 'companies', 'jobTitles', 'experience', 'experienceMonths', 'skills', 'summary']}, ensure_ascii=False)}

Nama file:
{filename}

TEKS CV:
{text[:14000]}
""".strip()


def call_ollama_json(prompt: str, timeout: int | None = None) -> dict[str, Any] | None:
    model = get_ollama_model()
    request_timeout = timeout or int(get_config("OLLAMA_TIMEOUT", "120") or "120")
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "Return only valid JSON. Do not add markdown, prose, or code fences.",
            },
            {"role": "user", "content": prompt},
        ],
        "stream": False,
        "format": "json",
        "options": {
            "temperature": 0.1,
            "num_ctx": int(get_config("OLLAMA_NUM_CTX", "8192") or "8192"),
        },
    }
    request = urllib.request.Request(
        f"{get_ollama_base_url()}/api/chat",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urlopen_direct(request, timeout=request_timeout) as response:  # noqa: S310
            response_payload = json.loads(response.read().decode("utf-8"))
    except urllib.error.HTTPError as exc:
        try:
            error_payload = json.loads(exc.read().decode("utf-8", errors="replace"))
            message = error_payload.get("error") or f"HTTP {exc.code}"
        except Exception:  # noqa: BLE001
            message = f"HTTP {exc.code}"
        raise RuntimeError(message) from exc
    content = ""
    if isinstance(response_payload.get("message"), dict):
        content = str(response_payload["message"].get("content") or "")
    if not content:
        content = str(response_payload.get("response") or "")
    parsed = parse_json_string(content)
    if isinstance(parsed, dict):
        parsed.setdefault("model", model)
        return parsed
    return None


def parse_resume_locally(filename: str, text: str, extraction_meta: dict[str, Any]) -> dict[str, Any]:
    normalized = normalize_text(text)
    lines = normalize_lines(text)
    work_zone_lines = extract_work_zone(lines)

    email = find_email(text)
    name_result = find_name(lines, filename, email)
    phone = find_phone(text)
    experience_result = find_experience(lines, normalized, work_zone_lines)
    skills = find_skills(normalized)
    age = find_age(normalized)
    education_result = find_education(lines)
    companies = find_companies(lines, normalized, work_zone_lines)
    job_titles = find_job_titles(work_zone_lines)

    evidence = []
    evidence.extend(name_result["evidence"])
    evidence.extend(experience_result["evidence"])
    evidence.extend(education_result["evidence"])
    if companies:
        evidence.append({"field": "companies", "value": " | ".join(companies[:4]), "source": "work history / company pattern", "confidence": 0.72})
    if skills:
        evidence.append({"field": "skills", "value": ", ".join(skills[:8]), "source": "keyword evidence", "confidence": 0.7})

    field_confidence = {
        "name": name_result["confidence"],
        "experience": experience_result["confidence"],
        "education": education_result["confidence"],
        "companies": 0.75 if companies else 0.25,
        "skills": 0.72 if skills else 0.25,
    }
    confidence = round(sum(field_confidence.values()) / len(field_confidence), 2)

    warnings: list[str] = []
    if extraction_meta.get("warning"):
        warnings.append(str(extraction_meta["warning"]))
    if name_result["confidence"] < 0.65:
        warnings.append("Nama kandidat confidence rendah; sistem memakai kandidat terbaik dari CV/filename.")
    if experience_result["years"] == 0:
        warnings.append("Total pengalaman tidak eksplisit terbaca; scoring memakai 0 tahun sampai ada bukti tanggal/pengalaman.")
    elif experience_result.get("rangeCount", 0) == 0:
        warnings.append("Pengalaman terbaca dari klaim total, bukan dari rentang tanggal kerja.")
    if len(normalized) < 500:
        warnings.append("Teks CV relatif pendek; kemungkinan layout/scan membuat sebagian data tidak terbaca.")

    return {
        "fileName": filename,
        "name": name_result["value"],
        "email": email,
        "phone": phone,
        "age": age,
        "education": education_result["value"],
        "companies": companies,
        "jobTitles": job_titles,
        "experience": experience_result["years"],
        "experienceMonths": experience_result["months"],
        "rawExperience": experience_result["label"],
        "skills": skills,
        "summary": build_summary(name_result["value"], email, phone, age, education_result["value"], experience_result["years"], skills),
        "cvText": normalized[:12000],
        "parserMode": "Local Smart Parser",
        "parseConfidence": confidence,
        "fieldConfidence": field_confidence,
        "parseEvidence": evidence[:10],
        "parseWarnings": warnings,
        "extractionMeta": extraction_meta,
        "localIntelligence": {
            "ruleVersion": "local-max-1",
            "dateRangeCount": experience_result.get("rangeCount", 0),
            "dateRanges": experience_result.get("ranges", []),
            "textLength": len(normalized),
            "workZoneLineCount": len(work_zone_lines),
        },
    }


def parse_resume_with_ai(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    raw: bytes = b"",
    extension: str = "",
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    api_key = get_openai_api_key()
    if not api_key:
        return None

    prompt = f"""
You are an expert Indonesian recruitment resume parser. Extract facts only from the CV text.
Do not invent missing values. If unsure, return null/empty and lower confidence.
Prefer the candidate's actual legal/display name, not job title, section heading, company name, or file label.
Calculate total work experience from date ranges when possible and state evidence.

Filename: {filename}
Local parser draft: {json.dumps({k: local_candidate.get(k) for k in ['name', 'email', 'phone', 'education', 'companies', 'experience', 'skills']}, ensure_ascii=False)}

CV TEXT:
{text[:24000]}
""".strip()

    schema = {
        "type": "object",
        "additionalProperties": False,
        "properties": {
            "name": {"type": ["string", "null"]},
            "email": {"type": ["string", "null"]},
            "phone": {"type": ["string", "null"]},
            "age": {"type": ["integer", "null"]},
            "education": {"type": ["string", "null"]},
            "companies": {"type": "array", "items": {"type": "string"}},
            "jobTitles": {"type": "array", "items": {"type": "string"}},
            "experience": {"type": "integer", "minimum": 0, "maximum": 50},
            "experienceMonths": {"type": "integer", "minimum": 0, "maximum": 600},
            "skills": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"},
            "confidence": {"type": "number", "minimum": 0, "maximum": 1},
            "fieldConfidence": {"type": "object", "additionalProperties": {"type": "number"}},
            "evidence": {
                "type": "array",
                "items": {
                    "type": "object",
                    "additionalProperties": False,
                    "properties": {
                        "field": {"type": "string"},
                        "value": {"type": "string"},
                        "source": {"type": "string"},
                        "confidence": {"type": "number", "minimum": 0, "maximum": 1},
                    },
                    "required": ["field", "value", "source", "confidence"],
                },
            },
            "warnings": {"type": "array", "items": {"type": "string"}},
        },
        "required": [
            "name",
            "email",
            "phone",
            "age",
            "education",
            "companies",
            "jobTitles",
            "experience",
            "experienceMonths",
            "skills",
            "summary",
            "confidence",
            "fieldConfidence",
            "evidence",
            "warnings",
        ],
    }

    user_content: list[dict[str, Any]] = [{"type": "input_text", "text": prompt}]
    if extension == ".pdf" and raw:
        user_content.append(
            {
                "type": "input_file",
                "filename": filename,
                "file_data": f"data:application/pdf;base64,{base64.b64encode(raw).decode('ascii')}",
            }
        )

    for model in get_resume_models():
        payload = {
            "model": model,
            "input": [
                {
                    "role": "system",
                    "content": [
                        {
                            "type": "input_text",
                            "text": (
                                "You are a conservative resume parsing engine for Indonesian HR systems. "
                                "Return only structured JSON matching the schema. Never invent facts. "
                                "Use the attached PDF visual/layout when present. If a field is uncertain, "
                                "lower its field confidence and add a warning."
                            ),
                        }
                    ],
                },
                {"role": "user", "content": user_content},
            ],
            "text": {"format": {"type": "json_schema", "name": "resume_parse", "schema": schema, "strict": False}},
            "max_output_tokens": 3200,
        }

        request = urllib.request.Request(
            "https://api.openai.com/v1/responses",
            data=json.dumps(payload).encode("utf-8"),
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            method="POST",
        )

        try:
            with urlopen_direct(request, timeout=75) as response:  # noqa: S310
                response_payload = json.loads(response.read().decode("utf-8"))
            parsed = extract_json_from_openai_response(response_payload)
            if isinstance(parsed, dict):
                parsed["parserMode"] = f"OpenAI Structured Parser â€¢ {model}"
                parsed.setdefault("warnings", [])
                return parsed
        except urllib.error.HTTPError as exc:
            errors.append(format_openai_http_error(model, exc) if errors is not None else "")
            continue
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
            if errors is not None:
                errors.append(f"{model}: {type(exc).__name__} - {str(exc)[:160]}.")
            continue
    return None


def format_openai_http_error(model: str, exc: urllib.error.HTTPError) -> str:
    try:
        payload = json.loads(exc.read().decode("utf-8", errors="replace"))
        error = payload.get("error", {})
        code = error.get("code") or error.get("type") or f"HTTP {exc.code}"
        message = error.get("message") or "OpenAI request failed"
        if code == "insufficient_quota":
            message = "quota/billing OpenAI belum aktif atau sudah habis"
        return f"{model}: {code} - {message}."
    except Exception:  # noqa: BLE001
        return f"{model}: HTTP {exc.code}."


def parse_resume_with_together(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    if not get_together_api_key():
        return None
    if len(text.strip()) < 40:
        if errors is not None:
            errors.append("Together: teks CV terlalu pendek untuk parser text-only.")
        return None

    prompt = build_resume_prompt(filename, text, local_candidate)
    try:
        parsed = call_together_json(prompt, gemini_resume_schema(), "resume_parse", temperature=0.1, timeout=75)
        if isinstance(parsed, dict):
            parsed["parserMode"] = f"Together Structured Parser â€¢ {parsed.get('model', get_together_model())}"
            parsed.setdefault("warnings", [])
            return parsed
    except Exception as exc:  # noqa: BLE001
        if errors is not None:
            errors.append(f"Together {get_together_model()}: {str(exc)[:180]}.")
    return None


def parse_resume_with_openrouter(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    if not get_openrouter_api_key():
        return None
    if len(text.strip()) < 40:
        if errors is not None:
            errors.append("OpenRouter: teks CV terlalu pendek untuk parser text-only.")
        return None

    prompt = build_resume_prompt(filename, text, local_candidate)
    try:
        parsed = call_openrouter_json(prompt, gemini_resume_schema(), "resume_parse", temperature=0.1, timeout=75)
        if isinstance(parsed, dict):
            parsed["parserMode"] = f"OpenRouter Structured Parser • {parsed.get('model', get_openrouter_model())}"
            parsed.setdefault("warnings", [])
            return parsed
    except Exception as exc:  # noqa: BLE001
        if errors is not None:
            errors.append(f"OpenRouter {get_openrouter_model()}: {str(exc)[:180]}.")
    return None


def parse_resume_with_gemini(
    filename: str,
    text: str,
    local_candidate: dict[str, Any],
    raw: bytes = b"",
    extension: str = "",
    errors: list[str] | None = None,
) -> dict[str, Any] | None:
    api_key = get_gemini_api_key()
    if not api_key:
        return None

    model = get_gemini_resume_model()
    prompt = build_resume_prompt(filename, text, local_candidate)
    parts: list[dict[str, Any]] = [{"text": prompt}]
    if extension == ".pdf" and raw:
        parts.append({"inline_data": {"mime_type": "application/pdf", "data": base64.b64encode(raw).decode("ascii")}})

    payload = {
        "contents": [{"role": "user", "parts": parts}],
        "generationConfig": {
            "temperature": 0.1,
            "responseMimeType": "application/json",
            "responseJsonSchema": gemini_resume_schema(),
        },
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    try:
        with urlopen_direct(request, timeout=75) as response:  # noqa: S310
            response_payload = json.loads(response.read().decode("utf-8"))
        parsed = extract_json_from_gemini_response(response_payload)
        if isinstance(parsed, dict):
            parsed["parserMode"] = f"Gemini Structured Parser â€¢ {model}"
            parsed.setdefault("warnings", [])
            return parsed
    except urllib.error.HTTPError as exc:
        if errors is not None:
            errors.append(format_gemini_http_error(model, exc))
    except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
        if errors is not None:
            errors.append(f"{model}: {type(exc).__name__} - {str(exc)[:160]}.")
    return None


def call_gemini_json(prompt: str, schema: dict[str, Any], temperature: float = 0.2, timeout: int = 75) -> dict[str, Any] | None:
    api_key = get_gemini_api_key()
    if not api_key:
        return None

    model = get_gemini_resume_model()
    payload = {
        "contents": [{"role": "user", "parts": [{"text": prompt}]}],
        "generationConfig": {
            "temperature": temperature,
            "responseMimeType": "application/json",
            "responseJsonSchema": schema,
        },
    }
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={api_key}"
    request = urllib.request.Request(
        url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json"},
        method="POST",
    )

    last_error: Exception | None = None
    for attempt in range(3):
        try:
            with urlopen_direct(request, timeout=timeout) as response:  # noqa: S310
                response_payload = json.loads(response.read().decode("utf-8"))
            parsed = extract_json_from_gemini_response(response_payload)
            if isinstance(parsed, dict):
                parsed.setdefault("model", model)
                return parsed
            return None
        except urllib.error.HTTPError as exc:
            last_error = exc
            if exc.code not in {429, 500, 502, 503, 504} or attempt == 2:
                raise RuntimeError(format_gemini_http_error(model, exc)) from exc
            time.sleep(1.5 * (attempt + 1))
        except urllib.error.URLError as exc:
            last_error = exc
            if attempt == 2:
                raise RuntimeError(f"{model}: network error - {str(exc)[:160]}.") from exc
            time.sleep(1.5 * (attempt + 1))
    if last_error:
        raise RuntimeError(str(last_error))
    return None


def call_openai_json(prompt: str, schema: dict[str, Any], schema_name: str, timeout: int = 75) -> dict[str, Any] | None:
    api_key = get_openai_api_key()
    if not api_key:
        return None

    last_error: Exception | None = None
    for model in get_resume_models():
        payload = {
            "model": model,
            "input": [
                {
                    "role": "system",
                    "content": [
                        {
                            "type": "input_text",
                            "text": "Return only valid JSON matching the provided schema. Do not add markdown.",
                        }
                    ],
                },
                {"role": "user", "content": [{"type": "input_text", "text": prompt}]},
            ],
            "text": {"format": {"type": "json_schema", "name": schema_name, "schema": schema, "strict": False}},
            "max_output_tokens": 3600,
        }
        request = urllib.request.Request(
            "https://api.openai.com/v1/responses",
            data=json.dumps(payload).encode("utf-8"),
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            method="POST",
        )
        try:
            with urlopen_direct(request, timeout=timeout) as response:  # noqa: S310
                response_payload = json.loads(response.read().decode("utf-8"))
            parsed = extract_json_from_openai_response(response_payload)
            if isinstance(parsed, dict):
                parsed.setdefault("model", model)
                return parsed
        except urllib.error.HTTPError as exc:
            last_error = exc
            code = exc.code
            if code not in {429, 500, 502, 503, 504}:
                raise RuntimeError(format_openai_http_error(model, exc)) from exc
            time.sleep(1)
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            continue
    if last_error:
        raise RuntimeError(str(last_error))
    return None


def call_together_json(
    prompt: str,
    schema: dict[str, Any],
    schema_name: str,
    temperature: float = 0.2,
    timeout: int = 75,
) -> dict[str, Any] | None:
    api_key = get_together_api_key()
    if not api_key:
        return None

    model = get_together_model()
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "Return only valid JSON matching the schema. Do not add markdown or commentary.",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
        "max_tokens": 3600,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": schema_name,
                "schema": schema,
            },
        },
    }
    request = urllib.request.Request(
        f"{get_together_base_url()}/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        method="POST",
    )

    last_error: Exception | None = None
    for attempt in range(3):
        try:
            with urlopen_direct(request, timeout=timeout) as response:  # noqa: S310
                response_payload = json.loads(response.read().decode("utf-8"))
            parsed = extract_json_from_chat_completions_response(response_payload)
            if isinstance(parsed, dict):
                parsed.setdefault("model", model)
                return parsed
            return None
        except urllib.error.HTTPError as exc:
            last_error = exc
            if exc.code not in {429, 500, 502, 503, 504} or attempt == 2:
                raise RuntimeError(format_openai_compatible_http_error("Together", model, exc)) from exc
            time.sleep(1.5 * (attempt + 1))
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            if attempt == 2:
                raise RuntimeError(f"Together {model}: {type(exc).__name__} - {str(exc)[:180]}.") from exc
            time.sleep(1.5 * (attempt + 1))
    if last_error:
        raise RuntimeError(str(last_error))
    return None


def call_openrouter_json(
    prompt: str,
    schema: dict[str, Any],
    schema_name: str,
    temperature: float = 0.2,
    timeout: int = 75,
) -> dict[str, Any] | None:
    api_key = get_openrouter_api_key()
    if not api_key:
        return None

    model = get_openrouter_model()
    payload = {
        "model": model,
        "messages": [
            {
                "role": "system",
                "content": "Return only valid JSON matching the schema. Do not add markdown or commentary.",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": temperature,
        "max_tokens": 3600,
        "response_format": {
            "type": "json_schema",
            "json_schema": {
                "name": schema_name,
                "schema": schema,
            },
        },
    }
    request = urllib.request.Request(
        f"{get_openrouter_base_url()}/chat/completions",
        data=json.dumps(payload).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
            "HTTP-Referer": "http://127.0.0.1:8765",
            "X-Title": "TalentFlow Intelligence",
        },
        method="POST",
    )

    last_error: Exception | None = None
    for attempt in range(3):
        try:
            with urlopen_direct(request, timeout=timeout) as response:  # noqa: S310
                response_payload = json.loads(response.read().decode("utf-8"))
            parsed = extract_json_from_chat_completions_response(response_payload)
            if isinstance(parsed, dict):
                parsed.setdefault("model", model)
                return parsed
            return None
        except urllib.error.HTTPError as exc:
            last_error = exc
            if exc.code not in {429, 500, 502, 503, 504} or attempt == 2:
                raise RuntimeError(format_openai_compatible_http_error("OpenRouter", model, exc)) from exc
            time.sleep(1.5 * (attempt + 1))
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError, ValueError) as exc:
            last_error = exc
            if attempt == 2:
                raise RuntimeError(f"OpenRouter {model}: {type(exc).__name__} - {str(exc)[:180]}.") from exc
            time.sleep(1.5 * (attempt + 1))
    if last_error:
        raise RuntimeError(str(last_error))
    return None


def call_multi_engine_json(prompt: str, schema: dict[str, Any], schema_name: str, temperature: float = 0.2, timeout: int = 75) -> dict[str, Any] | None:
    errors: list[str] = []
    for engine in select_ai_engines():
        try:
            if engine == "gemini":
                parsed = call_gemini_json(prompt, schema, temperature=temperature, timeout=timeout)
            elif engine == "openai":
                parsed = call_openai_json(prompt, schema, schema_name=schema_name, timeout=timeout)
            elif engine == "together":
                parsed = call_together_json(prompt, schema, schema_name=schema_name, temperature=temperature, timeout=timeout)
            elif engine == "openrouter":
                parsed = call_openrouter_json(prompt, schema, schema_name=schema_name, temperature=temperature, timeout=timeout)
            else:
                parsed = None
            if parsed:
                parsed["engine"] = engine
                return parsed
        except Exception as exc:  # noqa: BLE001
            errors.append(f"{engine}: {str(exc)[:220]}")
            continue
    if errors:
        raise RuntimeError(" | ".join(errors))
    return None


def generate_blueprint_with_ai(brief: dict[str, Any]) -> dict[str, Any] | None:
    prompt = f"""
You are a senior Indonesian Human Capital job analyst.
Create a practical Job Blueprint from the user's brief.
Do not write a job advertisement. Create the real working map of the role.
The output must use three editable sections only:
1. Tanggung Jawab Utama
2. Tantangan Jabatan
3. Persyaratan Jabatan

Use professional inference from position, level, industry, role, KPI orientation, and preferences.
Be specific to the level and industry. Avoid generic repetition.
Return JSON only.

USER BRIEF:
{json.dumps(brief, ensure_ascii=False)}
""".strip()
    parsed = call_multi_engine_json(prompt, gemini_blueprint_schema(), "job_blueprint", temperature=0.28)
    if not parsed:
        return None
    parsed["parserMode"] = f"{title_engine(parsed.get('engine'))} Job Blueprint Builder â€¢ {parsed.get('model', get_gemini_resume_model())}"
    return parsed


def score_candidate_fit_with_ai(candidate: dict[str, Any], blueprint: dict[str, Any]) -> dict[str, Any] | None:
    candidate_payload = {
        "name": candidate.get("name"),
        "role": candidate.get("role"),
        "age": candidate.get("age"),
        "education": candidate.get("education"),
        "experience": candidate.get("experience"),
        "rawExperience": candidate.get("rawExperience"),
        "companies": candidate.get("companies"),
        "jobTitles": candidate.get("jobTitles"),
        "skills": candidate.get("skills"),
        "summary": candidate.get("summary"),
        "cvText": str(candidate.get("cvText") or "")[:18000],
    }
    blueprint_payload = {
        "positionName": blueprint.get("positionName"),
        "companyName": blueprint.get("companyName"),
        "businessType": blueprint.get("businessType"),
        "roleHolder": blueprint.get("roleHolder"),
        "positionLevel": blueprint.get("positionLevel"),
        "preferences": blueprint.get("preferences"),
        "mainResponsibilities": blueprint.get("mainResponsibilities"),
        "roleChallenges": blueprint.get("roleChallenges"),
        "jobRequirements": blueprint.get("jobRequirements"),
    }
    prompt = f"""
You are a senior Indonesian recruiter and Human Capital assessor.
Score candidate fit against the FINAL Job Blueprint.

Language rule:
- All user-facing output values must be in Bahasa Indonesia.
- This includes status, reasons, inferences.reason, gaps, and interviewQuestions.
- Use professional, concise Indonesian. Avoid English terms unless they are common HR terms.

Important scoring rule:
- Do not depend on exact keyword matching.
- Read the candidate's career pattern, seniority, industry exposure, role scope, and likely capability.
- If a capability is not written explicitly but is strongly implied by job history, mark it as "Strong inference", not "Not found".
- Do not invent facts. Separate direct evidence from professional inference.
- Include risks and validation questions for interview.
- Use these Indonesian evidenceLevel labels when appropriate: "Bukti langsung", "Inferensi kuat", "Inferensi lemah", "Tidak ditemukan".
- Use Indonesian status labels such as "Sangat Sesuai", "Cukup Sesuai", "Potensial", or "Berisiko".

Weights:
- Tanggung Jawab Utama Fit: 40%
- Tantangan Jabatan Fit: 30%
- Persyaratan Jabatan Fit: 30%

Score scale rule:
- score and all componentScores must be integer 0-100, not 0-10.
- "Sangat Sesuai" should normally be 85-95.
- "Cukup Sesuai" should normally be 70-84.
- "Potensial" should normally be 55-69.
- "Berisiko" should normally be below 55.

Return JSON only.

FINAL JOB BLUEPRINT:
{json.dumps(blueprint_payload, ensure_ascii=False)}

CANDIDATE PROFILE AND CV TEXT:
{json.dumps(candidate_payload, ensure_ascii=False)}
""".strip()
    parsed = call_multi_engine_json(prompt, gemini_fit_score_schema(), "blueprint_fit_score", temperature=0.12, timeout=90)
    if not parsed:
        return None
    normalize_fit_score_scale(parsed)
    parsed["parserMode"] = f"{title_engine(parsed.get('engine'))} Blueprint Fit Scoring â€¢ {parsed.get('model', get_gemini_resume_model())}"
    return parsed


def title_engine(engine: Any) -> str:
    value = str(engine or "").lower()
    if value == "openai":
        return "OpenAI"
    if value == "gemini":
        return "Gemini"
    if value == "together":
        return "Together"
    if value == "openrouter":
        return "OpenRouter"
    return "AI"


def normalize_fit_score_scale(score_payload: dict[str, Any]) -> None:
    def normalize_score(value: Any) -> int:
        try:
            number = float(value)
        except (TypeError, ValueError):
            return 0
        if 0 < number <= 10:
            number *= 10
        return max(0, min(round(number), 100))

    score_payload["score"] = normalize_score(score_payload.get("score"))
    components = score_payload.get("componentScores")
    if isinstance(components, dict):
        for key in ("responsibilityFit", "challengeFit", "requirementFit"):
            components[key] = normalize_score(components.get(key))


def build_resume_prompt(filename: str, text: str, local_candidate: dict[str, Any]) -> str:
    return f"""
You are an expert Indonesian recruitment resume parser.
Extract facts only from the CV. Do not invent missing values.
The candidate name must be a person name, not a job title, company name, section heading, or filename label.
Calculate total work experience from employment date ranges when possible.
If the CV contains overlapping jobs, count overlapping months once.
Return JSON only.

Filename: {filename}
Local parser draft: {json.dumps({k: local_candidate.get(k) for k in ['name', 'email', 'phone', 'education', 'companies', 'experience', 'skills']}, ensure_ascii=False)}

CV TEXT:
{text[:24000]}
""".strip()


def gemini_resume_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "properties": {
            "name": {"type": "string"},
            "email": {"type": "string"},
            "phone": {"type": "string"},
            "age": {"type": "integer"},
            "education": {"type": "string"},
            "companies": {"type": "array", "items": {"type": "string"}},
            "jobTitles": {"type": "array", "items": {"type": "string"}},
            "experience": {"type": "integer"},
            "experienceMonths": {"type": "integer"},
            "skills": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string"},
            "confidence": {"type": "number"},
            "fieldConfidence": {
                "type": "object",
                "properties": {
                    "name": {"type": "number"},
                    "experience": {"type": "number"},
                    "education": {"type": "number"},
                    "companies": {"type": "number"},
                    "skills": {"type": "number"},
                },
            },
            "evidence": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "field": {"type": "string"},
                        "value": {"type": "string"},
                        "source": {"type": "string"},
                        "confidence": {"type": "number"},
                    },
                    "required": ["field", "value", "source", "confidence"],
                },
            },
            "warnings": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["name", "email", "phone", "education", "companies", "jobTitles", "experience", "experienceMonths", "skills", "summary", "confidence", "fieldConfidence", "evidence", "warnings"],
    }


def gemini_blueprint_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "properties": {
            "mainResponsibilities": {"type": "array", "items": {"type": "string"}},
            "roleChallenges": {"type": "array", "items": {"type": "string"}},
            "jobRequirements": {"type": "array", "items": {"type": "string"}},
            "scoringGuidance": {"type": "array", "items": {"type": "string"}},
            "confidence": {"type": "number"},
            "warnings": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["mainResponsibilities", "roleChallenges", "jobRequirements", "scoringGuidance", "confidence", "warnings"],
    }


def gemini_fit_score_schema() -> dict[str, Any]:
    return {
        "type": "object",
        "properties": {
            "score": {"type": "integer"},
            "status": {"type": "string"},
            "confidence": {"type": "number"},
            "componentScores": {
                "type": "object",
                "properties": {
                    "responsibilityFit": {"type": "integer"},
                    "challengeFit": {"type": "integer"},
                    "requirementFit": {"type": "integer"},
                },
                "required": ["responsibilityFit", "challengeFit", "requirementFit"],
            },
            "reasons": {"type": "array", "items": {"type": "string"}},
            "inferences": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "capability": {"type": "string"},
                        "evidenceLevel": {"type": "string"},
                        "reason": {"type": "string"},
                        "confidence": {"type": "number"},
                    },
                    "required": ["capability", "evidenceLevel", "reason", "confidence"],
                },
            },
            "gaps": {"type": "array", "items": {"type": "string"}},
            "interviewQuestions": {"type": "array", "items": {"type": "string"}},
        },
        "required": ["score", "status", "confidence", "componentScores", "reasons", "inferences", "gaps", "interviewQuestions"],
    }


def extract_json_from_gemini_response(payload: Any) -> dict[str, Any] | None:
    try:
        candidates = payload.get("candidates", [])
        for candidate in candidates:
            parts = candidate.get("content", {}).get("parts", [])
            for part in parts:
                text = part.get("text")
                if isinstance(text, str):
                    parsed = parse_json_string(text)
                    if isinstance(parsed, dict):
                        return parsed
    except AttributeError:
        return None
    return None


def format_gemini_http_error(model: str, exc: urllib.error.HTTPError) -> str:
    try:
        payload = json.loads(exc.read().decode("utf-8", errors="replace"))
        error = payload.get("error", {})
        code = error.get("status") or error.get("code") or f"HTTP {exc.code}"
        message = error.get("message") or "Gemini request failed"
        return f"{model}: {code} - {message[:180]}."
    except Exception:  # noqa: BLE001
        return f"{model}: HTTP {exc.code}."


def extract_json_from_openai_response(payload: Any) -> dict[str, Any] | None:
    if isinstance(payload, dict):
        for key in ("output_text", "text"):
            value = payload.get(key)
            if isinstance(value, str):
                parsed = parse_json_string(value)
                if isinstance(parsed, dict):
                    return parsed
        for value in payload.values():
            parsed = extract_json_from_openai_response(value)
            if parsed:
                return parsed
    elif isinstance(payload, list):
        for item in payload:
            parsed = extract_json_from_openai_response(item)
            if parsed:
                return parsed
    return None


def extract_json_from_chat_completions_response(payload: Any) -> dict[str, Any] | None:
    try:
        choices = payload.get("choices", [])
        for choice in choices:
            message = choice.get("message", {})
            content = message.get("content")
            if isinstance(content, str):
                parsed = parse_json_string(content)
                if isinstance(parsed, dict):
                    return parsed
            if isinstance(content, list):
                for item in content:
                    text = item.get("text") if isinstance(item, dict) else None
                    if isinstance(text, str):
                        parsed = parse_json_string(text)
                        if isinstance(parsed, dict):
                            return parsed
    except AttributeError:
        return None
    return None


def format_openai_compatible_http_error(provider: str, model: str, exc: urllib.error.HTTPError) -> str:
    try:
        payload = json.loads(exc.read().decode("utf-8", errors="replace"))
        error = payload.get("error", {})
        code = error.get("code") or error.get("type") or f"HTTP {exc.code}"
        message = error.get("message") or f"{provider} request failed"
        return f"{provider} {model}: {code} - {message[:220]}."
    except Exception:  # noqa: BLE001
        return f"{provider} {model}: HTTP {exc.code}."


def parse_json_string(value: str) -> Any:
    try:
        return json.loads(value)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", value, re.DOTALL)
        if match:
            return json.loads(match.group(0))
    return None


def merge_ai_candidate(local_candidate: dict[str, Any], ai_candidate: dict[str, Any]) -> dict[str, Any]:
    merged = dict(local_candidate)
    ai_confidence = float(ai_candidate.get("confidence") or 0)
    field_confidence = {**local_candidate.get("fieldConfidence", {}), **(ai_candidate.get("fieldConfidence") or {})}

    def choose_text(field: str) -> str:
        ai_value = clean_scalar(ai_candidate.get(field))
        local_value = clean_scalar(local_candidate.get(field))
        if ai_value and float(field_confidence.get(field, ai_confidence)) >= 0.45:
            return ai_value
        return local_value

    merged["name"] = choose_text("name") or local_candidate.get("name")
    merged["email"] = choose_text("email") or local_candidate.get("email")
    merged["phone"] = choose_text("phone") or local_candidate.get("phone")
    merged["education"] = choose_text("education") or local_candidate.get("education")
    merged["summary"] = choose_text("summary") or local_candidate.get("summary")
    merged["age"] = ai_candidate.get("age") or local_candidate.get("age")
    merged["companies"] = dedupe([*safe_companies(ai_candidate), *safe_string_list(local_candidate.get("companies"))])[:8]
    merged["jobTitles"] = dedupe([*safe_job_titles(ai_candidate), *safe_string_list(local_candidate.get("jobTitles"))])[:8]
    merged["skills"] = dedupe([*safe_string_list(ai_candidate.get("skills")), *safe_string_list(local_candidate.get("skills"))])[:16]

    ai_months = safe_months(ai_candidate.get("experienceMonths"))
    local_months = safe_months(local_candidate.get("experienceMonths"))
    months = max(ai_months, local_months)
    ai_years = safe_int(ai_candidate.get("experience"))
    local_years = safe_int(local_candidate.get("experience"))
    merged["experienceMonths"] = months
    merged["experience"] = max(ai_years, local_years, round(months / 12))
    merged["rawExperience"] = months_to_label(months) if months else local_candidate.get("rawExperience", "")
    merged["parserMode"] = ai_candidate.get("parserMode", "OpenAI Structured Parser")
    merged["parseConfidence"] = round(max(ai_confidence, float(local_candidate.get("parseConfidence") or 0)), 2)
    merged["fieldConfidence"] = field_confidence
    merged["parseEvidence"] = [*(ai_candidate.get("evidence") or []), *(local_candidate.get("parseEvidence") or [])][:12]
    merged["parseWarnings"] = dedupe([*(ai_candidate.get("warnings") or []), *(local_candidate.get("parseWarnings") or [])])[:8]
    return merged


def clean_scalar(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, dict):
        preferred = [
            value.get("name"),
            value.get("fullName"),
            value.get("education"),
            " ".join(str(part) for part in (value.get("degreeLevel"), value.get("major"), value.get("schoolName")) if part),
            value.get("companyName"),
            value.get("positionTitle"),
            value.get("value"),
        ]
        for item in preferred:
            cleaned = clean_scalar(item)
            if cleaned:
                return cleaned
        return ""
    return re.sub(r"\s+", " ", str(value)).strip()


def safe_int(value: Any) -> int:
    if isinstance(value, bool) or value is None:
        return 0
    if isinstance(value, (int, float)):
        return max(0, round(value))
    if isinstance(value, dict):
        for key in ("months", "totalMonths", "value", "years", "totalYears"):
            if key in value:
                number = safe_int(value.get(key))
                if number:
                    return number
        return 0
    match = re.search(r"\d+(?:[,.]\d+)?", str(value))
    if not match:
        return 0
    number = float(match.group(0).replace(",", "."))
    return max(0, round(number))


def safe_months(value: Any) -> int:
    if isinstance(value, dict):
        for key in ("months", "totalMonths", "value"):
            number = safe_int(value.get(key))
            if number:
                return number
        for key in ("years", "totalYears"):
            number = safe_int(value.get(key))
            if number:
                return number * 12
        return 0
    text = str(value or "").lower()
    number = safe_int(value)
    if number and re.search(r"\b(year|years|tahun|thn)\b", text):
        return number * 12
    return number


def safe_string_list(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    output: list[str] = []
    for item in value:
        if isinstance(item, dict):
            if item.get("companyName"):
                output.append(clean_scalar(item.get("companyName")))
            elif item.get("positionTitle"):
                output.append(clean_scalar(item.get("positionTitle")))
            elif item.get("skill"):
                output.append(clean_scalar(item.get("skill")))
            else:
                output.append(clean_scalar(item))
        else:
            output.append(clean_scalar(item))
    return [item for item in output if item]


def safe_companies(candidate: dict[str, Any]) -> list[str]:
    companies = safe_string_list(candidate.get("companies"))
    if companies:
        return companies
    history = candidate.get("workExperience") or candidate.get("employmentHistory") or candidate.get("experiences")
    if isinstance(history, list):
        return [clean_scalar(item.get("companyName") or item.get("company") or item.get("employer")) for item in history if isinstance(item, dict)]
    return []


def safe_job_titles(candidate: dict[str, Any]) -> list[str]:
    titles = safe_string_list(candidate.get("jobTitles"))
    if titles:
        return titles
    history = candidate.get("workExperience") or candidate.get("employmentHistory") or candidate.get("experiences")
    if isinstance(history, list):
        return [clean_scalar(item.get("positionTitle") or item.get("title") or item.get("role")) for item in history if isinstance(item, dict)]
    return []


def normalize_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.lower()).strip()


def normalize_lines(text: str) -> list[str]:
    return [re.sub(r"\s+", " ", line).strip() for line in text.splitlines() if re.sub(r"\s+", " ", line).strip()]


def find_email(text: str) -> str:
    match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text, re.IGNORECASE)
    return match.group(0) if match else ""


def find_phone(text: str) -> str:
    compact = re.sub(r"(?<!\w)[+]?62", "+62", text)
    match = re.search(r"(\+62|62|0)\s*8(?:[\s().-]?\d){7,12}", compact)
    if match:
        return re.sub(r"\s+", " ", match.group(0)).strip()
    return ""


def find_experience(lines: list[str], normalized: str, work_zone_lines: list[str]) -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    explicit_months = find_explicit_experience_months(normalized)
    if explicit_months:
        evidence.append({"field": "experience", "value": months_to_label(explicit_months), "source": "total pengalaman eksplisit", "confidence": 0.88})

    range_detail = calculate_date_range_experience_detail(work_zone_lines or lines)
    range_months = int(range_detail["months"])
    if range_months:
        evidence.append(
            {
                "field": "experience",
                "value": f"{months_to_label(range_months)} dari {len(range_detail['ranges'])} periode kerja",
                "source": "rentang tanggal pengalaman kerja",
                "confidence": 0.84 if len(range_detail["ranges"]) >= 2 else 0.78,
            }
        )

    months = max(explicit_months, range_months)
    if explicit_months and range_months:
        confidence = 0.92
    elif explicit_months:
        confidence = 0.88
    elif range_months:
        confidence = 0.84 if len(range_detail["ranges"]) >= 2 else 0.78
    else:
        confidence = 0.2
    return {
        "years": round(months / 12) if months else 0,
        "months": months,
        "label": months_to_label(months) if months else "",
        "confidence": confidence,
        "evidence": evidence,
        "rangeCount": len(range_detail["ranges"]),
        "ranges": range_detail["ranges"][:8],
    }


def find_explicit_experience_months(text: str) -> int:
    patterns = [
        r"(?:total\s*)?(?:pengalaman|experience|experiences|working experience|work experience)[^\d]{0,35}(\d{1,2})(?:[,.](\d))?\+?\s*(?:tahun|thn|years?|yrs?)",
        r"(\d{1,2})(?:[,.](\d))?\+?\s*(?:tahun|thn|years?|yrs?)\s*(?:pengalaman|experience|experiences)",
        r"berpengalaman[^\d]{0,25}(\d{1,2})(?:[,.](\d))?\+?\s*(?:tahun|thn|years?|yrs?)",
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            context = text[max(0, match.start() - 24) : match.end() + 12]
            if re.search(r"(usia|umur|age)\s*[:\-]?\s*$", context[: match.start() - max(0, match.start() - 24)], re.IGNORECASE):
                continue
            years = int(match.group(1))
            decimal = match.group(2)
            months = years * 12 + (int(decimal) * 12 // 10 if decimal else 0)
            return min(months, 600)
    return 0


def calculate_date_range_experience(lines: list[str]) -> int:
    return int(calculate_date_range_experience_detail(lines)["months"])


def calculate_date_range_experience_detail(lines: list[str]) -> dict[str, Any]:
    month_name = r"jan(?:uary|uari)?|feb(?:ruary|ruari)?|mar(?:ch|et)?|apr(?:il)?|may|mei|jun(?:e|i)?|jul(?:y|i)?|aug(?:ust|ustus)?|sep(?:t|tember)?|oct(?:ober)?|okt(?:ober)?|nov(?:ember)?|dec(?:ember)?|des(?:ember)?"
    textual_date = rf"(?:(?:{month_name})[\s./-]*)?(?:19|20)\d{{2}}"
    numeric_date = r"(?:0?[1-9]|1[0-2])[\s./-](?:19|20)\d{2}"
    present_date = r"present|current|now|sekarang|saat ini|hingga kini|sampai sekarang"
    date_ref = rf"(?:{numeric_date}|{textual_date}|{present_date})"
    separator = r"(?:-|\u2013|\u2014|to|until|sampai|hingga|sd|s/d)"
    range_pattern = re.compile(rf"(?P<start>{date_ref})\s*{separator}\s*(?P<end>{date_ref})", re.IGNORECASE)

    covered_months: set[int] = set()
    seen_ranges: set[tuple[int, int]] = set()
    ranges: list[dict[str, Any]] = []
    zone = "\n".join(lines)
    for match in range_pattern.finditer(zone):
        start_ref = match.group("start").strip()
        end_ref = match.group("end").strip()
        start_date = parse_date_ref(start_ref, is_end=False)
        end_date = parse_date_ref(end_ref, is_end=True)
        if not start_date or not end_date:
            continue
        start_index = start_date[0] * 12 + start_date[1]
        end_index = end_date[0] * 12 + end_date[1]
        if end_index < start_index or end_index - start_index > 600:
            continue
        key = (start_index, end_index)
        if key in seen_ranges:
            continue
        seen_ranges.add(key)
        for month_index in range(start_index, end_index + 1):
            covered_months.add(month_index)
        ranges.append({"start": start_ref, "end": end_ref, "months": end_index - start_index + 1})
    return {"months": min(len(covered_months), 600), "ranges": ranges}


def parse_date_ref(value: str, is_end: bool) -> tuple[int, int] | None:
    cleaned = value.lower().strip()
    if cleaned in {"present", "current", "now", "sekarang", "saat ini", "hingga kini", "sampai sekarang"}:
        return CURRENT_YEAR, CURRENT_MONTH
    numeric_match = re.search(r"\b(0?[1-9]|1[0-2])[\s./-]((?:19|20)\d{2})\b", cleaned)
    if numeric_match:
        year = int(numeric_match.group(2))
        month = int(numeric_match.group(1))
        if 1970 <= year <= CURRENT_YEAR:
            return year, month
    year_match = re.search(r"(19|20)\d{2}", cleaned)
    if not year_match:
        return None
    year = int(year_match.group(0))
    month = 12 if is_end else 1
    for token, month_number in MONTHS.items():
        if re.search(rf"\b{re.escape(token)}\b", cleaned):
            month = month_number
            break
    if year < 1970 or year > CURRENT_YEAR:
        return None
    return year, month


def months_to_label(months: int) -> str:
    if months <= 0:
        return ""
    years, remaining_months = divmod(months, 12)
    if years and remaining_months:
        return f"{years} tahun {remaining_months} bulan"
    if years:
        return f"{years} tahun"
    return f"{remaining_months} bulan"

def find_age(text: str) -> int | None:
    patterns = [r"(?:usia|umur|age)\s*[:\-]?\s*(\d{2})", r"(\d{2})\s*(?:tahun|years old|yo)\b"]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            age = int(match.group(1))
            if 17 <= age <= 65:
                return age
    birth_year_match = re.search(r"(?:lahir|birth|born|tanggal lahir)[^\d]*(19\d{2}|20\d{2})", text, re.IGNORECASE)
    if birth_year_match:
        age = CURRENT_YEAR - int(birth_year_match.group(1))
        if 17 <= age <= 65:
            return age
    return None


def find_education(lines: list[str]) -> dict[str, Any]:
    rank_terms = [
        ("s3", 6), ("doctor", 6), ("phd", 6),
        ("s2", 5), ("master", 5), ("magister", 5),
        ("s1", 4), ("sarjana", 4), ("bachelor", 4),
        ("d4", 3), ("d3", 3), ("diploma", 3),
        ("smk", 2), ("sma", 2),
    ]
    school_terms = ("universitas", "university", "institut", "institute", "politeknik", "academy", "akademi", "sekolah tinggi")
    major_terms = ("jurusan", "major", "program studi", "fakultas", "teknik", "manajemen", "psikologi", "hukum", "akuntansi", "ekonomi", "informatika", "industri")
    best = {"value": "", "confidence": 0.2, "rank": 0, "evidence": []}
    for index, line in enumerate(lines):
        cleaned = re.sub(r"\s+", " ", line).strip(" |?-")
        if not cleaned or len(cleaned) > 160:
            continue
        lowered = cleaned.lower()
        context_lines = [cleaned]
        for neighbor in (index - 1, index + 1):
            if 0 <= neighbor < len(lines):
                nearby = re.sub(r"\s+", " ", lines[neighbor]).strip(" |?-")
                if 0 < len(nearby) <= 120 and any(term in nearby.lower() for term in (*school_terms, *major_terms)):
                    context_lines.append(nearby)
        combined = " | ".join(dedupe(context_lines))
        combined_lower = combined.lower()
        rank = 0
        for term, score in rank_terms:
            if re.search(rf"\b{re.escape(term)}\b", combined_lower):
                rank = max(rank, score)
        has_school = any(term in combined_lower for term in school_terms)
        has_major = any(term in combined_lower for term in major_terms)
        if rank or has_school:
            confidence = 0.68 + (0.1 if has_school else 0) + (0.06 if has_major else 0) + (0.04 if rank else 0)
            if rank > best["rank"] or (rank == best["rank"] and confidence >= best["confidence"]):
                best = {
                    "value": combined,
                    "confidence": min(confidence, 0.92),
                    "rank": rank,
                    "evidence": [{"field": "education", "value": combined, "source": "baris pendidikan dan konteks sekitar", "confidence": min(confidence, 0.92)}],
                }
    return best

def extract_work_zone(lines: list[str]) -> list[str]:
    start_markers = ("pengalaman", "work experience", "professional experience", "employment", "riwayat kerja", "career history")
    end_markers = ("education", "pendidikan", "skill", "keahlian", "certification", "sertifikasi", "training")
    start = None
    for index, line in enumerate(lines):
        lowered = line.lower()
        if any(marker in lowered for marker in start_markers):
            start = index
            break
    if start is None:
        return lines
    end = len(lines)
    for index in range(start + 1, len(lines)):
        lowered = lines[index].lower()
        if any(marker in lowered for marker in end_markers):
            end = index
            break
    return lines[start:end]


def find_companies(lines: list[str], normalized: str, work_zone_lines: list[str]) -> list[str]:
    companies: list[str] = []
    company_pattern = re.compile(r"\b(?:pt\.?|cv\.?|ud\.?|tbk\.?|ltd\.?|inc\.?|corp\.?|company|group)\s+[a-z0-9&.,'() -]{2,55}", re.IGNORECASE)

    for line in work_zone_lines:
        cleaned = re.sub(r"\s+", " ", line).strip(" |â€¢-")
        lowered = cleaned.lower()
        if len(cleaned) > 120 or looks_like_section(cleaned):
            continue
        company_match = company_pattern.search(cleaned)
        if company_match:
            companies.append(title_company(strip_date_tail(company_match.group(0))))
            continue
        if has_date_range(cleaned):
            continue
        if re.search(r"\s[-â€“â€”]\s", cleaned):
            possible_company = re.split(r"\s[-â€“â€”]\s", cleaned)[-1].strip()
            if possible_company and not looks_like_job_title(possible_company) and not has_date_range(possible_company):
                companies.append(title_company(possible_company))
                continue
        if re.search(r"\b(di|at)\s+[A-Z][A-Za-z0-9&.,'() -]{2,55}", cleaned):
            candidate = re.sub(r"^.*?\b(?:di|at)\s+", "", cleaned, flags=re.IGNORECASE)
            companies.append(title_company(candidate))
            continue
    return dedupe([company for company in companies if is_plausible_company(company)])[:8]


def title_company(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" |â€¢-")
    replacements = {"pt": "PT", "cv": "CV", "tbk": "Tbk", "ltd": "Ltd", "inc": "Inc"}
    parts = []
    for token in value.split():
        normalized = token.strip(".").lower()
        parts.append(replacements.get(normalized, token[:1].upper() + token[1:]))
    return " ".join(parts)


def strip_date_tail(value: str) -> str:
    return re.sub(
        r"\b(?:jan(?:uary|uari)?|feb(?:ruary|ruari)?|mar(?:ch|et)?|apr(?:il)?|may|mei|jun(?:e|i)?|jul(?:y|i)?|aug(?:ust|ustus)?|sep(?:t|tember)?|oct(?:ober)?|okt(?:ober)?|nov(?:ember)?|dec(?:ember)?|des(?:ember)?|(?:19|20)\d{2}).*$",
        "",
        value,
        flags=re.IGNORECASE,
    ).strip(" -â€“â€”|,")


def is_plausible_company(value: str) -> bool:
    lowered = value.lower()
    if len(value) < 3 or len(value) > 100:
        return False
    if any(word in lowered for word in ("email", "phone", "skill", "curriculum", "resume")):
        return False
    return True


def find_job_titles(work_zone_lines: list[str]) -> list[str]:
    titles: list[str] = []
    title_terms = ("manager", "supervisor", "specialist", "officer", "staff", "lead", "head", "coordinator", "analyst", "partner", "recruiter", "hrbp", "director", "assistant", "kepala", "koordinator", "admin", "operator", "gudang", "warehouse", "hrd", "general affairs", "ga")
    description_starts = ("mengelola", "memimpin", "bertanggung", "melakukan", "membuat", "menyusun", "menangani", "support", "responsible", "manage ", "managed ", "handling")
    for line in work_zone_lines:
        cleaned = re.sub(r"\s+", " ", line).strip(" |?-")
        lowered = cleaned.lower()
        has_title_term = any(re.search(rf"\b{re.escape(term)}\b", lowered) for term in title_terms)
        if not has_title_term:
            continue
        if any(lowered.startswith(prefix) for prefix in description_starts):
            continue
        if "," in cleaned and len(cleaned) > 45:
            continue
        if 3 <= len(cleaned) <= 90 and not looks_like_section(cleaned):
            cleaned = re.sub(r"\(?\b(?:19|20)\d{2}\b.*$", "", cleaned).strip(" -??|,")
            if re.search(r"\s[-??]\s", cleaned):
                cleaned = re.split(r"\s[-??]\s", cleaned)[0].strip()
            if cleaned:
                titles.append(cleaned)
    return dedupe(titles)[:8]

def find_skills(text: str) -> list[str]:
    found = [skill for skill in SKILL_KEYWORDS if re.search(rf"\b{re.escape(skill)}\b", text, re.IGNORECASE)]
    return found[:18]


def find_name(lines: list[str], filename: str, email: str = "") -> dict[str, Any]:
    evidence: list[dict[str, Any]] = []
    explicit_pattern = re.compile(r"^(?:nama lengkap|full name|nama|name)\s*[:\-|]\s*(?P<name>.+)$", re.IGNORECASE)
    for line in lines[:30]:
        match = explicit_pattern.search(line)
        if match:
            cleaned = clean_name_candidate(match.group("name"))
            if is_plausible_name(cleaned):
                evidence.append({"field": "name", "value": cleaned, "source": "explicit name label", "confidence": 0.94})
                return {"value": cleaned, "confidence": 0.94, "evidence": evidence}

    scored_candidates: list[tuple[float, str, str]] = []
    for index, line in enumerate(lines[:28]):
        for candidate in split_name_candidates(line):
            cleaned = clean_name_candidate(candidate)
            if not is_plausible_name(cleaned):
                continue
            score = score_name_candidate(cleaned, index, filename, email)
            if score >= 0.45:
                scored_candidates.append((score, cleaned, f"line {index + 1}: {line[:90]}"))

    if scored_candidates:
        scored_candidates.sort(key=lambda item: item[0], reverse=True)
        score, value, source = scored_candidates[0]
        confidence = min(round(score, 2), 0.92)
        evidence.append({"field": "name", "value": value, "source": source, "confidence": confidence})
        return {"value": value, "confidence": confidence, "evidence": evidence}

    fallback = filename_to_name(filename)
    evidence.append({"field": "name", "value": fallback, "source": "filename fallback", "confidence": 0.45})
    return {"value": fallback, "confidence": 0.45, "evidence": evidence}


def split_name_candidates(line: str) -> list[str]:
    cleaned = re.sub(r"[|â€¢]+", "\n", line)
    parts = [part.strip() for part in re.split(r"\n| {3,}", cleaned) if part.strip()]
    return parts or [line]


def clean_name_candidate(value: str) -> str:
    value = re.sub(r"\b(curriculum vitae|resume|cv|profile|profil)\b", "", value, flags=re.IGNORECASE)
    value = re.sub(r"\b(bapak|ibu|mr|mrs|ms|dr|prof)\.?\s+", "", value, flags=re.IGNORECASE)
    value = re.sub(r"https?://\S+|www\.\S+", "", value, flags=re.IGNORECASE)
    value = re.sub(r"[,;].*$", "", value)
    value = re.sub(r"\s+", " ", value)
    return value.strip(" -â€“â€”|:/")


def is_plausible_name(value: str) -> bool:
    if not value or "@" in value or re.search(r"\d", value):
        return False
    lowered = value.lower().strip()
    if lowered in SECTION_WORDS or any(word in lowered for word in ("linkedin", "portfolio", "github", "address", "alamat", "phone", "email", "mobile", "whatsapp")):
        return False
    if re.search(r"\b(pt|cv|tbk|ltd|inc|corp|company|universitas|university|sekolah|institute|institut)\b", lowered):
        return False
    words = value.split()
    if not 2 <= len(words) <= 5:
        return False
    if len(value) > 60:
        return False
    alpha_words = [word for word in words if re.search(r"[A-Za-z]", word)]
    if len(alpha_words) != len(words):
        return False
    job_hits = sum(1 for word in words if word.lower().strip(".,") in JOB_WORDS)
    if job_hits >= max(2, len(words) - 1):
        return False
    return True


def score_name_candidate(value: str, index: int, filename: str, email: str = "") -> float:
    words = value.split()
    score = 0.35
    if index <= 6:
        score += 0.24
    elif index <= 14:
        score += 0.12
    if 2 <= len(words) <= 3:
        score += 0.16
    if value.isupper() or sum(1 for word in words if word[:1].isupper()) >= len(words) - 1:
        score += 0.13
    if not any(word.lower() in JOB_WORDS for word in words):
        score += 0.08
    filename_tokens = set(re.findall(r"[a-z]+", filename.lower())) - {"cv", "resume", "lamaran", "new", "update", "final"}
    value_tokens = set(re.findall(r"[a-z]+", value.lower()))
    if filename_tokens and filename_tokens & value_tokens:
        score += 0.12
    email_tokens = set(re.findall(r"[a-z]+", email.split("@", 1)[0].lower())) - {"hr", "recruitment", "admin"}
    if email_tokens and value_tokens and len(email_tokens & value_tokens) >= 1:
        score += 0.12
    if index > 18:
        score -= 0.08
    return score


def filename_to_name(filename: str) -> str:
    stem = Path(filename).stem.lower()
    stem = re.sub(r"\b(cv|resume|curriculum|vitae|lamaran|application|final|updated?|new|baru)\b", " ", stem)
    stem = re.sub(r"\b(?:19|20)\d{2}\b|\d+", " ", stem)
    stem = re.sub(r"[_\-.]+", " ", stem)
    words = [word for word in stem.split() if word not in JOB_WORDS and len(word) > 1]
    if not words:
        words = [word for word in re.sub(r"[_\-.]+", " ", Path(filename).stem).split() if len(word) > 1]
    return " ".join(word.capitalize() for word in words[:4]) or "Kandidat Tanpa Nama"


def has_date_range(line: str) -> bool:
    return bool(
        re.search(
            r"(?:19|20)\d{2}.*(?:-|â€“|â€”|to|until|sampai|hingga|sd|s/d).*(?:(?:19|20)\d{2}|present|current|sekarang|saat ini|hingga kini)",
            line,
            re.IGNORECASE,
        )
    )


def looks_like_section(line: str) -> bool:
    lowered = line.lower().strip(" :")
    return lowered in SECTION_WORDS or (len(lowered.split()) <= 3 and any(word in lowered for word in SECTION_WORDS))


def looks_like_job_title(line: str) -> bool:
    lowered_words = set(re.findall(r"[a-z]+", line.lower()))
    return bool(lowered_words & JOB_WORDS) and not re.search(r"\b(pt|cv|ltd|inc|tbk)\b", line.lower())


def dedupe(items: list[str]) -> list[str]:
    seen = set()
    output = []
    for item in items:
        cleaned = re.sub(r"\s+", " ", str(item)).strip()
        normalized = cleaned.lower()
        if not cleaned or normalized in seen:
            continue
        seen.add(normalized)
        output.append(cleaned)
    return output


def build_summary(name: str, email: str, phone: str, age: int | None, education: str, experience: int, skills: list[str]) -> str:
    parts = [name]
    if age:
        parts.append(f"{age} tahun")
    if education:
        parts.append(education)
    if experience:
        parts.append(f"{experience} tahun pengalaman")
    if skills:
        parts.append(", ".join(skills[:4]))
    if email:
        parts.append(email)
    if phone:
        parts.append(phone)
    return " | ".join(parts)


def run() -> None:
    mimetypes.add_type("text/javascript", ".js")
    host = os.getenv("HOST") or ("0.0.0.0" if os.getenv("PORT") else "127.0.0.1")
    server = ThreadingHTTPServer((host, PORT), RecruitFlowHandler)
    display_host = "127.0.0.1" if host == "0.0.0.0" else host
    print(f"RecruitFlow running at http://{display_host}:{PORT}")
    server.serve_forever()


if __name__ == "__main__":
    run()
